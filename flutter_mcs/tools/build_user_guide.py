from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\Somya Raj\MCS-Hitachi\docs\MCS_User_Guide.docx")
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
RED = "E60012"
INK = "1A1D23"
MUTED = "64748B"
LIGHT = "E8EEF5"
PALE_RED = "FFF1F2"
GREEN = "166534"
GOLD = "8A5A00"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc.get_or_add_tcPr()
    tc_mar = tc.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_cell(cell, bold=False, color=INK, size=9.5):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.1
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        header.cells[idx].text = text
        shade(header.cells[idx], LIGHT)
        format_cell(header.cells[idx], bold=True, color=DARK_BLUE)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
            format_cell(cells[idx])
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_page_field(paragraph, field_name):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, end))


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(.375)
        style.paragraph_format.first_line_indent = Inches(-.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = RGBColor.from_string(INK)
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.left_indent = Inches(.18)
    callout.paragraph_format.right_indent = Inches(.18)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(.492)
    section.footer_distance = Inches(.492)
    header = section.header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("MCS USER GUIDE")
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(MUTED)
    p.add_run("  ·  Merchant Checkout System").font.color.rgb = RGBColor.from_string(MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    rr = fp.add_run("Hitachi Payments BRD implementation  ·  Page ")
    rr.font.name = "Calibri"
    rr.font.size = Pt(9)
    rr.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(fp, "PAGE")


def title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(INK)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(18)
        rr = p2.add_run(subtitle)
        rr.font.name = "Calibri"
        rr.font.size = Pt(13)
        rr.font.color.rgb = RGBColor.from_string(MUTED)


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(doc, text):
    doc.add_paragraph(text, style="List Number")


def callout(doc, heading, text, fill=PALE_RED, compact=False):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    row_pr = table.rows[0]._tr.get_or_add_trPr()
    row_pr.append(OxmlElement("w:cantSplit"))
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_after = Pt(0 if compact else 3)
    r = p.add_run(heading)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(RED)
    if compact:
        p.add_run("  " + text)
    else:
        p.paragraph_format.keep_with_next = True
        p2 = cell.add_paragraph(text)
        p2.paragraph_format.keep_together = True
        p2.paragraph_format.space_after = Pt(0)
    set_table_width(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build():
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    core = doc.core_properties
    core.title = "MCS Merchant and Customer User Guide"
    core.subject = "Hitachi Payments BRD operating guide"
    core.author = "MCS Project Team"
    core.keywords = "MCS, merchant, customer, billing, payment, refund"

    # Editorial-cover first page: quiet running furniture, generous whitespace.
    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(14)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run("OPERATIONS & USER GUIDE")
    kr.font.name = "Calibri"
    kr.font.size = Pt(11)
    kr.font.bold = True
    kr.font.color.rgb = RGBColor.from_string(RED)
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(10)
    cover.paragraph_format.space_after = Pt(8)
    cr = cover.add_run("MCS")
    cr.font.name = "Calibri"
    cr.font.size = Pt(42)
    cr.font.bold = True
    cr.font.color.rgb = RGBColor.from_string(RED)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Merchant Checkout System")
    sr.font.name = "Calibri"
    sr.font.size = Pt(22)
    sr.font.bold = True
    sr.font.color.rgb = RGBColor.from_string(INK)
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.paragraph_format.space_after = Pt(55)
    tr = tagline.add_run("A practical guide for merchants, customers, administrators, and demo reviewers")
    tr.font.name = "Calibri"
    tr.font.size = Pt(12)
    tr.font.italic = True
    tr.font.color.rgb = RGBColor.from_string(MUTED)
    add_table(doc, ["Edition", "Scope", "Status"], [["1.0 · 14 July 2026", "BRD end-to-end processes", "Verified build"]], [2200, 4560, 2600])
    callout(doc, "What this guide covers", "Account access, merchant billing, customer payment, transaction states, refunds, reports, live activity, assistant behavior, and recovery steps.")

    page_break(doc)
    title(doc, "1. Start here", "Roles, services, sign-in, and the safest first run")
    doc.add_heading("System roles", level=1)
    add_table(doc, ["Role", "Primary work", "Protected boundary"], [
        ["Merchant", "Create/find bills, monitor customers and settlements, refund paid bills, export reports", "Only the merchant's own bills, customers, transactions, and reports"],
        ["Customer", "View bills, pay pending bills, review payment history", "Only the customer's own bills and transactions"],
    ], [1700, 4300, 3360])
    doc.add_heading("First local run", level=1)
    numbered(doc, "Start PostgreSQL and confirm the mcs_db database exists.")
    numbered(doc, "Start Kafka on port 9092 when demonstrating the live feed.")
    numbered(doc, "Run the Spring Boot API on port 8080.")
    numbered(doc, "Run Flutter with MCS_API_BASE_URL pointing to the API.")
    numbered(doc, "Sign in as the intended role; merchant and customer sessions are deliberately separate.")
    doc.add_heading("Seeded demo access", level=2)
    add_table(doc, ["Role", "Email", "Password"], [
        ["Merchant", "sharma@electronics.com", "McsDemo@123"],
        ["Customer", "rahul@gmail.com", "McsDemo@123"],
    ], [1800, 4300, 3260])
    callout(doc, "Security", "A 401 means the token is absent or expired. A 403 means the signed-in account does not own the requested resource. Do not work around either response by changing IDs.", "FFF7ED")

    page_break(doc)
    title(doc, "2. Merchant workflow", "From bill creation to customer visibility and settlement reporting")
    doc.add_heading("Create a bill", level=1)
    numbered(doc, "Open Bills and choose Create Bill.")
    numbered(doc, "Select the intended customer and enter an amount of at least ₹1.00.")
    numbered(doc, "Add a description that lets both parties identify the charge.")
    numbered(doc, "Submit once. The new bill appears as PENDING.")
    doc.add_heading("Find a bill without long scrolling", level=1)
    bullet(doc, "Search by bill ID, customer, or description.")
    bullet(doc, "Filter by Pending, Paid, Failed, or Refunded; combine with the date filter.")
    bullet(doc, "Sort newest/oldest as needed. Clear filters to restore the full directory.")
    doc.add_heading("Connected customers", level=1)
    doc.add_paragraph("The Customers section is live backend data aggregated from the merchant's persisted bills. A customer becomes connected when the merchant has bill activity with that customer. Paid value and outstanding value are calculated from bill status, not typed into the dashboard.")
    doc.add_heading("Live transaction feed", level=1)
    bullet(doc, "Settled payments publish a Kafka transaction-completed event.")
    bullet(doc, "The dashboard receives the merchant's activity through the WebSocket feed.")
    bullet(doc, "Use the horizontal/contained scrolling area for additional cards; it is intentionally not a tall page list.")

    page_break(doc)
    title(doc, "3. Customer workflow", "Review, pay, verify, and recover")
    doc.add_heading("Understand the dashboard", level=1)
    add_table(doc, ["Area", "Meaning"], [
        ["Outstanding", "Total value of currently pending bills"],
        ["Bills", "Actual bill count and status distribution"],
        ["Spending trend", "Settled payment values by real activity date; labels/tooltips expose amounts"],
        ["Recent activity", "Latest persisted bills/payments, not mock examples"],
        ["Next payment", "The nearest actionable pending bill, or an all-caught-up state"],
    ], [2300, 7060])
    doc.add_heading("Pay a bill", level=1)
    numbered(doc, "Open Bills, filter to Pending, and select the correct bill.")
    numbered(doc, "Review merchant, description, and amount before continuing.")
    numbered(doc, "Choose UPI, Card, or Net Banking and initiate the transaction.")
    numbered(doc, "The API returns an authorization decision. A decline leaves the bill pending for a safe retry.")
    numbered(doc, "After authorization, settlement changes the bill to PAID and creates a unique reference.")
    doc.add_heading("Verify the result", level=1)
    bullet(doc, "The success receipt must show Transaction ID, settlement reference, merchant, amount, and PAID status.")
    bullet(doc, "Payment History can be searched and filtered; use it instead of relying only on the dashboard card.")
    callout(doc, "Never double-pay", "If the screen is interrupted after authorization, reopen the bill and inspect its status/history before attempting another payment.")

    page_break(doc)
    title(doc, "4. Lifecycle, refunds, and reports", "Status meanings and operator rules")
    doc.add_heading("Transaction states", level=1)
    add_table(doc, ["State", "Meaning", "Next valid action"], [
        ["INITIATED", "Payment attempt recorded", "Authorize"],
        ["AUTHORIZED", "Payment method approved", "Settle"],
        ["SETTLED", "Funds recorded and settlement persisted", "Report or refund the paid bill"],
        ["FAILED", "Authorization declined", "Retry from the pending bill"],
    ], [1900, 4300, 3160])
    doc.add_heading("Refund a payment", level=1)
    numbered(doc, "Merchant opens a bill with PAID status.")
    numbered(doc, "Choose Refund and enter a concise reason.")
    numbered(doc, "Confirm once. The backend records REFUNDED, the reason, and the refund timestamp.")
    numbered(doc, "Download the refund CSV from Reports when evidence is required.")
    callout(doc, "Refund rule", "Pending and failed bills cannot be refunded because no successful settlement exists.", "FFF7ED")
    doc.add_heading("Report outputs", level=1)
    add_table(doc, ["Output", "Contents", "Use"], [
        ["Daily report", "Today's settlement totals and counts", "Daily reconciliation"],
        ["Weekly report", "Seven-day settlement totals and counts", "Operational review"],
        ["Bill statement CSV", "Merchant bill directory", "Spreadsheet analysis/audit"],
        ["Refund CSV", "Only persisted REFUNDED bills", "Refund control evidence"],
    ], [2050, 4200, 3110])

    page_break(doc)
    title(doc, "5. Assistant and operational recovery", "What is safe, what is live, and what to do when a service is down")
    doc.add_heading("Assistant behavior", level=1)
    bullet(doc, "The assistant greets the signed-in merchant or customer once per login session.")
    bullet(doc, "It summarizes account activity through role-aware backend logic and optional classifier support.")
    bullet(doc, "It does not claim to create bills inside chat unless a real bill-creation action is implemented and confirmed.")
    bullet(doc, "Clear removes the current conversation display; it does not delete bills, payments, or account data.")
    doc.add_heading("Recovery matrix", level=1)
    add_table(doc, ["Symptom", "Likely cause", "Action"], [
        ["Customer service unavailable / 404", "Old backend or wrong base URL", "Start the current API and verify /customers/merchant/{merchantId}"],
        ["Live feed stays still", "Kafka/WebSocket unavailable or no new settlements", "Start Kafka, reconnect, then settle a new test bill"],
        ["Refund report empty", "No REFUNDED bills", "Refund one PAID bill, then export again"],
        ["401", "Missing/expired JWT", "Sign in again"],
        ["403", "Wrong role or resource owner", "Use the correct account; do not substitute IDs"],
        ["Physical device cannot reach API", "localhost points at the device", "Pass the computer IP through MCS_API_BASE_URL"],
    ], [2450, 3230, 3680])
    doc.add_heading("Data integrity checks", level=1)
    bullet(doc, "One successful transaction must create exactly one settlement.")
    bullet(doc, "A failed authorization must not create a settlement or mark the bill paid.")
    bullet(doc, "A refund must preserve the original transaction/settlement evidence and update the bill to REFUNDED.")

    page_break(doc)
    title(doc, "6. Acceptance and support reference", "How to prove the build and hand it over")
    doc.add_heading("Automated evidence", level=1)
    add_table(doc, ["Check", "Command", "Expected"], [
        ["Backend suite", ".\\mvnw.cmd test", "Context, service, lifecycle, and security tests pass"],
        ["Flutter suite", "flutter test", "Payment result and role registration tests pass without overflow"],
        ["1,000 transactions", ".\\mvnw.cmd --% -q -Dtest=TransactionScalabilityTest -Dmcs.performance.tests=true test", "1,000 settled; 0 failures; p95 under 2 s"],
        ["Postman", "Run MCS BRD API Acceptance", "All executed requests under 2 s; no 5xx"],
    ], [2050, 4670, 2640])
    doc.add_heading("Measured scalability result", level=1)
    add_table(doc, ["Merchants", "Transactions", "REST calls", "Failures", "p95"], [["100", "1,000", "3,000", "0", "9.45 ms"]], [1800, 2000, 1900, 1500, 2160])
    doc.add_paragraph("Environment: Spring Boot application context, JWT/MockMvc REST boundary, JPA, and H2 PostgreSQL mode. Kafka/network capacity is excluded; repeat against deployed PostgreSQL/Kafka before making a production infrastructure claim.")
    doc.add_heading("Handover artifacts", level=1)
    bullet(doc, "README.md — setup and operating instructions")
    bullet(doc, "postman/ — chained API acceptance collection and local environment")
    bullet(doc, "docs/MCS_ER_Diagram.svg — database relationships")
    bullet(doc, "docs/performance/1000-transactions-report.md — measured benchmark")
    bullet(doc, "docs/MCS_Final_Demo.pptx — final presentation")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
